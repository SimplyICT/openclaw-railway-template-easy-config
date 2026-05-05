import os
import sys
from docx import Document
from docx.shared import Inches

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

def md_to_docx(md_filename):
    if not os.path.exists(md_filename):
        print(f"Error: {md_filename} not found.")
        return

    doc_filename = md_filename.replace('.md', '.docx')
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(md_filename, 'r') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph() # Add spacing
            continue
        
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Brand styling
            run = h.runs[0]
            run.font.color.rgb = RGBColor(211, 47, 47) # Red theme
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=1)
            run = h.runs[0]
            run.font.color.rgb = RGBColor(51, 51, 51)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=2)
            run = h.runs[0]
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(211, 47, 47)
        elif line.startswith('• ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            # Handle inline bold in bullets
            parts = line[2:].split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True

    doc.save(doc_filename)
    print(f"Successfully beautified {doc_filename}")
    return doc_filename

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 md_to_docx.py <input.md>")
    else:
        md_to_docx(sys.argv[1])
